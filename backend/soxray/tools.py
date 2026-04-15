import pandas as pd
from pathlib import Path
from typing import Any, cast
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

from soxray.models import EvidenceFile, TestFinding, WorkpaperOutput, ControlDefinition

_current_control: ControlDefinition | None = None
_findings_buffer: list[TestFinding] = []


def set_current_control(control: ControlDefinition) -> None:
    global _current_control, _findings_buffer
    _current_control = control
    _findings_buffer = []


def _record_finding(finding: TestFinding) -> None:
    _findings_buffer.append(finding)


def _get_current_context() -> tuple[ControlDefinition, list[TestFinding]]:
    if _current_control is None:
        raise ValueError("Current control is not set for workpaper generation.")
    return _current_control, list(_findings_buffer)


def load_evidence(filename: str) -> EvidenceFile:
    df = pd.read_csv(filename)
    parsed = df.replace({pd.NA: None}).to_dict(orient="records")
    
    evidence_type = "unknown"
    if "termination" in filename.lower():
        evidence_type = "HR_report"
    elif "ad_events" in filename.lower():
        evidence_type = "AD_log"
        
    return EvidenceFile(
        filename=filename,
        file_type="csv",
        parsed_content=parsed,
        evidence_type=evidence_type
    )

def join_datasets(df1: list[dict[str, Any]], df2: list[dict[str, Any]], join_key: str) -> list[dict[str, Any]]:
    pdf1 = pd.DataFrame(df1)
    pdf2 = pd.DataFrame(df2)
    
    if pdf1.empty or pdf2.empty:
        return []
        
    merged = pd.merge(pdf1, pdf2, on=join_key, how="left")
    return cast(list[dict[str, Any]], merged.replace({pd.NA: None}).to_dict(orient="records"))

def calculate_delta(timestamp1: Any, timestamp2: Any, unit: str = "hours") -> float:
    if pd.isna(timestamp1) or pd.isna(timestamp2) or not timestamp1 or not timestamp2:
        return -1.0
        
    t1 = pd.to_datetime(timestamp1)
    t2 = pd.to_datetime(timestamp2)
    
    delta = t2 - t1
    
    if unit == "hours":
        return delta.total_seconds() / 3600.0
    elif unit == "days":
        return delta.total_seconds() / (3600.0 * 24.0)
    return delta.total_seconds()

def lookup_record(dataset: list[dict[str, Any]], key: str, value: Any) -> dict[str, Any]:
    for record in dataset:
        if record.get(key) == value:
            return record
    
    return {"status": "not_found", "message": f"No record found with {key}={value}"}

def flag_exception(sample_id: str, sample_identifier: str, finding_detail: str, citations: list[str]) -> TestFinding:
    finding = TestFinding(
        sample_id=sample_id,
        sample_identifier=sample_identifier,
        result="EXCEPTION",
        evidence_citations=citations,
        finding_detail=finding_detail
    )
    _record_finding(finding)
    return finding

def flag_pass(sample_id: str, sample_identifier: str, citations: list[str]) -> TestFinding:
    finding = TestFinding(
        sample_id=sample_id,
        sample_identifier=sample_identifier,
        result="PASS",
        evidence_citations=citations,
        finding_detail="Sample passed testing without exceptions."
    )
    _record_finding(finding)
    return finding


def generate_workpaper_from_context(total_samples: int, exceptions: int, conclusion: str) -> str:
    control, findings = _get_current_context()

    actual_total = len(findings)
    total = total_samples if total_samples > 0 else actual_total

    actual_exceptions = sum(1 for f in findings if f.result == "EXCEPTION")
    exc_count = exceptions if exceptions > 0 else actual_exceptions

    output = WorkpaperOutput(
        control=control,
        total_samples=total,
        exceptions=exc_count,
        findings=findings,
        conclusion=conclusion,
    )
    return write_workpaper(output)

def write_workpaper(output: WorkpaperOutput) -> str:
    doc = Document()
    
    title = doc.add_heading(f"SOX Workpaper: {output.control.control_id}", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph(f"Control Name: {output.control.control_name}", style="Intense Quote")
    
    header_table = doc.add_table(rows=3, cols=2)
    header_table.style = 'Table Grid'
    hdr_cells = header_table.rows[0].cells
    hdr_cells[0].text = "Test Date"
    hdr_cells[1].text = datetime.now().strftime("%Y-%m-%d")
    
    hdr_cells = header_table.rows[1].cells
    hdr_cells[0].text = "Preparer"
    hdr_cells[1].text = "soxray.ai Automated Agent"
    
    hdr_cells = header_table.rows[2].cells
    hdr_cells[0].text = "Control Type / Freq"
    hdr_cells[1].text = f"{output.control.control_type} / {output.control.frequency}"
    
    doc.add_paragraph()
    
    doc.add_heading("Control Description", level=1)
    doc.add_paragraph(output.control.control_description)
    
    doc.add_heading("Test Procedure", level=1)
    doc.add_paragraph(output.control.test_procedure)
    
    doc.add_heading("Results Summary", level=1)
    summary_text = (f"Total Samples Tested: {output.total_samples}\n"
                    f"Total Exceptions Identified: {output.exceptions}\n"
                    f"Exception Rate: {(output.exceptions / max(1, output.total_samples)) * 100:.2f}%")
    doc.add_paragraph(summary_text)

    doc.add_heading("Sample Testing Selection & Results", level=1)
    sample_table = doc.add_table(rows=1, cols=4)
    sample_table.style = 'Light Shading Accent 1'
    head_cells = sample_table.rows[0].cells
    head_cells[0].text = "Sample ID"
    head_cells[1].text = "Identifier"
    head_cells[2].text = "Result"
    head_cells[3].text = "Evidence Citations"
    
    for finding in output.findings:
        row_cells = sample_table.add_row().cells
        row_cells[0].text = finding.sample_id
        row_cells[1].text = finding.sample_identifier
        row_cells[2].text = finding.result
        row_cells[3].text = ", ".join(finding.evidence_citations)

    if output.exceptions > 0:
        doc.add_heading("Exception Details", level=1)
        for finding in output.findings:
            if finding.result == "EXCEPTION":
                p = doc.add_paragraph()
                p.add_run(f"Sample {finding.sample_id} ({finding.sample_identifier}): ").bold = True
                p.add_run(finding.finding_detail)
                
                resp_p = doc.add_paragraph("Management Response: ")
                resp_p.runs[0].italic = True
                resp_p.add_run(finding.control_owner_response or "Pending Input")

    doc.add_heading("Conclusion", level=1)
    conclusion_p = doc.add_paragraph()
    conclusion_p.add_run(output.conclusion).bold = True
    
    out_dir = Path("output")
    out_dir.mkdir(exist_ok=True)
    filename = out_dir / f"workpaper_{output.control.control_id.lower()}.docx"
    doc.save(str(filename))
    return str(filename)
